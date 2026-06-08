import argparse
import os
import sys


def build_docx(title: str, body: str, out: str):
    try:
        from docx import Document
    except Exception as exc:
        raise SystemExit(f"python-docx not installed: {exc}")
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(body)
    doc.save(out)
    print(f"Wrote {out}")


def build_pdf(title: str, body: str, out: str):
    try:
        from fpdf import FPDF
    except Exception:
        try:
            from reportlab.pdfgen import canvas
        except Exception as exc:
            raise SystemExit(f"fpdf2/reportlab not installed: {exc}")
        c = canvas.Canvas(out)
        c.drawString(72, 760, title)
        text = c.beginText(72, 740)
        for line in body.splitlines() or [body]:
            text.textLine(line)
        c.drawText(text)
        c.save()
        print(f"Wrote {out}")
        return
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 20)
    pdf.cell(0, 10, title, ln=1)
    pdf.set_font("Helvetica", "", 12)
    pdf.multi_cell(0, 8, body)
    pdf.output(out)
    print(f"Wrote {out}")


def build_odt(title: str, body: str, out: str):
    try:
        from odf import opendocument, text
    except Exception as exc:
        raise SystemExit(f"odfpy not installed: {exc}")
    doc = opendocument.OpenDocumentText()
    doc.text.addElement(text.H(text=title))
    doc.text.addElement(text.P(text=body))
    doc.save(out)
    print(f"Wrote {out}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--title", required=True)
    parser.add_argument("--body", required=True)
    parser.add_argument("--out", required=True)
    args = parser.parse_args()
    ext = os.path.splitext(args.out)[1].lower().lstrip(".")
    if ext == "docx":
        build_docx(args.title, args.body, args.out)
    elif ext == "pdf":
        build_pdf(args.title, args.body, args.out)
    elif ext in {"odt", "odp", "ods"}:
        build_odt(args.title, args.body, args.out)
    else:
        raise SystemExit(f"Unsupported extension: .{ext}")


if __name__ == "__main__":
    main()
